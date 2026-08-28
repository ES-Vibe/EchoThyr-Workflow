"""
Word document generation module using win32com (COM automation)
Generates medical reports from template with bookmarks
Compatible with PowerShell version bookmark system
"""

from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass
from datetime import datetime


@dataclass
class PatientInfo:
    """Patient information"""
    last_name: str = "A PRECISER"
    first_name: str = ""
    exam_date: str = ""
    birth_date: str = ""  # Date de naissance (format DD.MM.YYYY)


class WordGenerator:
    """Generate Word documents from template using COM automation"""

    def __init__(self, template_path: str):
        self.template_path = template_path

    def extract_patient_info(self, folder_name: str) -> PatientInfo:
        """Extract patient info from folder name (format: 'NOM Prenom')"""
        parts = folder_name.split()
        info = PatientInfo(exam_date=datetime.now().strftime("%d.%m.%Y"))

        if len(parts) >= 1:
            info.last_name = parts[0].upper()
        if len(parts) >= 2:
            info.first_name = parts[1].capitalize()

        return info

    def generate_report(
        self,
        patient_info: PatientInfo,
        measurements: List,
        image_paths: List[str],
        output_path: str,
        logger=None
    ) -> bool:
        """
        Generate medical report from template using Word COM automation

        Args:
            patient_info: Patient information
            measurements: List of Measurement objects
            image_paths: List of resized image paths
            output_path: Path for output DOCX file
            logger: Optional logger

        Returns:
            True if successful, False otherwise
        """
        word = None
        doc = None

        try:
            import win32com.client
            import pythoncom

            # Initialize COM
            pythoncom.CoInitialize()

            # Start Word application
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Open template
            doc = word.Documents.Open(self.template_path)

            # Helper function to set bookmark or replace placeholder text
            def set_bookmark(name: str, content: str):
                if doc.Bookmarks.Exists(name):
                    bookmark_range = doc.Bookmarks.Item(name).Range
                    bookmark_range.Text = content
                    doc.Bookmarks.Add(name, bookmark_range)
                    if logger:
                        logger.debug(f"Bookmark set: {name} = {content[:50]}...")
                else:
                    # Fallback: replace placeholder text [NAME]
                    replace_placeholder(f"[{name}]", content)

            def replace_placeholder(placeholder: str, replacement: str):
                """Replace placeholder text in the document by scanning paragraphs."""
                found = False
                try:
                    for para in doc.Paragraphs:
                        para_range = para.Range
                        para_text = para_range.Text
                        if placeholder in para_text:
                            find_obj = para_range.Find
                            find_obj.ClearFormatting()
                            find_obj.Replacement.ClearFormatting()
                            result = find_obj.Execute(
                                FindText=placeholder,
                                ReplaceWith=replacement,
                                Replace=2,
                                Forward=True,
                                Wrap=0,
                                MatchCase=False
                            )
                            if result:
                                found = True
                    if not found:
                        for para in doc.Paragraphs:
                            para_text = para.Range.Text
                            if placeholder.lower() in para_text.lower():
                                import re
                                new_text = re.sub(re.escape(placeholder), replacement, para_text, flags=re.IGNORECASE)
                                if new_text != para_text:
                                    para.Range.Text = new_text
                                    found = True
                    if logger:
                        if found:
                            logger.debug(f"Placeholder replaced: {placeholder} -> {replacement[:50]}...")
                        else:
                            logger.warning(f"Placeholder not found in document: {placeholder}")
                except Exception as e:
                    if logger:
                        logger.warning(f"Error replacing placeholder {placeholder}: {e}")

            # Set patient info bookmarks
            set_bookmark("NOM", patient_info.last_name)
            set_bookmark("PRENOM", patient_info.first_name)
            set_bookmark("DATE", patient_info.exam_date)
            if patient_info.birth_date:
                set_bookmark("DATE_NAISSANCE", patient_info.birth_date)

            # Generate measurement text
            measurement_text = self._format_measurements(measurements)
            set_bookmark("RESULTAT", measurement_text)

            # Add images at the end
            if image_paths:
                # Go to end of document
                doc.Characters.Last.Select()
                word.Selection.InsertBreak(7)  # Page break

                for img_path in image_paths:
                    try:
                        word.Selection.InlineShapes.AddPicture(img_path)
                        word.Selection.TypeText("\r\n")
                    except Exception as e:
                        if logger:
                            logger.warning(f"Failed to add image {img_path}: {e}")

            # Save document (16 = wdFormatDocumentDefault)
            doc.SaveAs2(output_path, 16)

            if logger:
                logger.success(f"Word document generated: {output_path}")

            return True

        except ImportError as e:
            if logger:
                logger.error(f"win32com not available: {e}")
            return False

        except Exception as e:
            if logger:
                logger.error(f"Failed to generate Word document: {e}", exc_info=e)
            return False

        finally:
            # Cleanup COM objects
            try:
                if doc:
                    doc.Close(False)
                if word:
                    word.Quit()
            except:
                pass

    def generate_report_with_text(
        self,
        patient_info: PatientInfo,
        measurement_text: str,
        image_paths: List[str],
        output_path: str,
        logger=None,
        max_retries: int = 2,
        schema_path: str = None,
        nodule_table: dict = None
    ) -> bool:
        """
        Generate medical report with pre-formatted measurement text (from SR)

        Args:
            patient_info: Patient information
            measurement_text: Pre-formatted measurement text from SR
            image_paths: List of image paths to embed
            output_path: Path for output DOCX file
            logger: Optional logger
            max_retries: Number of retry attempts on failure
            schema_path: Optional path to thyroid schema PNG
            nodule_table: Optional measurement table (src.schema.build_table),
                inserted at the [TABLEAU] placeholder if the template has one

        Returns:
            True if successful, False otherwise
        """
        # Verify template exists first
        template_path = Path(self.template_path).resolve()
        if not template_path.exists():
            if logger:
                logger.error(f"Template file not found: {template_path}")
            return False

        last_error = None
        for attempt in range(max_retries + 1):
            if attempt > 0:
                if logger:
                    logger.info(f"Retry attempt {attempt}/{max_retries}...")
                # Wait a bit before retrying
                import time
                time.sleep(1)
                # Kill any hanging Word processes before retry
                self._kill_word_processes(logger)

            result = self._generate_report_internal(
                patient_info, measurement_text, image_paths,
                output_path, template_path, logger,
                schema_path=schema_path,
                nodule_table=nodule_table
            )
            if result:
                return True

        return False

    def _kill_word_processes(self, logger=None):
        """Kill any hanging Word processes"""
        try:
            import subprocess
            result = subprocess.run(
                ['taskkill', '/F', '/IM', 'WINWORD.EXE'],
                capture_output=True, text=True
            )
            if logger and 'SUCCESS' in result.stdout:
                logger.debug("Killed hanging Word process")
        except:
            pass

    def _replace_placeholders_docx(self, template_path: Path, output_path: str,
                                     patient_info: PatientInfo, measurement_text: str,
                                     logger=None) -> bool:
        """Replace placeholder text using python-docx (handles split runs reliably)"""
        try:
            from docx import Document

            doc = Document(str(template_path))

            replacements = {
                "[NOM]": patient_info.last_name,
                "[PRENOM]": patient_info.first_name,
                "[DATE]": patient_info.exam_date,
                "[RESULTAT]": measurement_text,
            }
            if patient_info.birth_date:
                replacements["[DATE_NAISSANCE]"] = patient_info.birth_date

            for paragraph in doc.paragraphs:
                full_text = paragraph.text
                for placeholder, value in replacements.items():
                    if placeholder in full_text:
                        # Rebuild runs: find which runs contain parts of the placeholder
                        self._replace_in_paragraph(paragraph, placeholder, value)
                        if logger:
                            logger.info(f"Replaced {placeholder} -> {value[:50]}")

            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            full_text = paragraph.text
                            for placeholder, value in replacements.items():
                                if placeholder in full_text:
                                    self._replace_in_paragraph(paragraph, placeholder, value)
                                    if logger:
                                        logger.info(f"Replaced {placeholder} in table -> {value[:50]}")

            doc.save(output_path)
            return True

        except Exception as e:
            if logger:
                logger.error(f"python-docx replacement failed: {e}", exc_info=e)
            return False

    def _replace_in_paragraph(self, paragraph, placeholder: str, replacement: str):
        """Replace placeholder in a paragraph, handling split runs and multiline text"""
        from docx.oxml.ns import qn

        runs = paragraph.runs
        full_text = "".join(run.text for run in runs)

        if placeholder not in full_text:
            return

        start_idx = full_text.index(placeholder)
        end_idx = start_idx + len(placeholder)

        # Find the run containing the start of the placeholder
        char_count = 0
        for i, run in enumerate(runs):
            run_start = char_count
            run_end = char_count + len(run.text)

            if run_start <= start_idx < run_end:
                before = run.text[:start_idx - run_start]

                # Clear subsequent runs that are part of the placeholder
                if end_idx > run_end:
                    for j in range(i + 1, len(runs)):
                        next_start = sum(len(runs[k].text) for k in range(j))
                        next_end = next_start + len(runs[j].text)
                        if next_end <= end_idx:
                            runs[j].text = ""
                        else:
                            runs[j].text = runs[j].text[end_idx - next_start:]
                            break
                    after = ""
                else:
                    after = run.text[end_idx - run_start:]

                # Handle multiline replacement with soft line breaks
                lines = replacement.replace("\r\n", "\n").split("\n")
                if len(lines) <= 1:
                    # Simple single-line replacement
                    run.text = before + replacement + after
                else:
                    # Multiline: first line in current run, then add breaks + runs
                    run.text = before + lines[0]
                    # Add remaining lines with line breaks (soft return)
                    for line_idx, line in enumerate(lines[1:], 1):
                        # Add a break element after the current run
                        br = run._element.makeelement(qn('w:br'), {})
                        run._element.append(br)
                        # Add text after the break using a text element
                        if line_idx < len(lines) - 1:
                            # Not the last line: add text then prepare for next break
                            t = run._element.makeelement(qn('w:t'), {})
                            t.text = line
                            t.set(qn('xml:space'), 'preserve')
                            run._element.append(t)
                        else:
                            # Last line: add text + after
                            t = run._element.makeelement(qn('w:t'), {})
                            t.text = line + after
                            t.set(qn('xml:space'), 'preserve')
                            run._element.append(t)
                return

            char_count = run_end

    def _generate_report_internal(
        self,
        patient_info: PatientInfo,
        measurement_text: str,
        image_paths: List[str],
        output_path: str,
        template_path: Path,
        logger=None,
        schema_path: str = None,
        nodule_table: dict = None
    ) -> bool:
        """Internal method to generate the report using python-docx only (no COM)"""
        try:
            from docx import Document
            from docx.shared import Inches

            # Step 1: Replace placeholders
            doc = Document(str(template_path))

            replacements = {
                "[NOM]": patient_info.last_name,
                "[PRENOM]": patient_info.first_name,
                "[DATE]": patient_info.exam_date,
                "[RESULTAT]": measurement_text,
            }
            if patient_info.birth_date:
                replacements["[DATE_NAISSANCE]"] = patient_info.birth_date

            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        self._replace_in_paragraph(paragraph, placeholder, value)
                        if logger:
                            logger.info(f"Replaced {placeholder} -> {value[:50]}")

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacements.items():
                                if placeholder in paragraph.text:
                                    self._replace_in_paragraph(paragraph, placeholder, value)

            # Step 1b: Insert thyroid schema if available
            schema_paragraph = None
            if schema_path and Path(schema_path).exists():
                # Try to find and replace [SCHEMA] placeholder
                for paragraph in doc.paragraphs:
                    if "[SCHEMA]" in paragraph.text:
                        # Clear the placeholder text
                        for run in paragraph.runs:
                            run.text = ""
                        # Insert schema image in this paragraph
                        run = paragraph.add_run()
                        run.add_picture(schema_path, width=Inches(5.0))
                        schema_paragraph = paragraph
                        if logger:
                            logger.info("Schema inserted at [SCHEMA] placeholder")
                        break

                if schema_paragraph is None:
                    # No placeholder found: insert schema after the last paragraph
                    # (before the image page break)
                    doc.add_paragraph("")  # Spacing
                    schema_paragraph = doc.add_paragraph()
                    run = schema_paragraph.add_run()
                    run.add_picture(schema_path, width=Inches(5.0))
                    if logger:
                        logger.info("Schema inserted at end of document (no [SCHEMA] placeholder found)")

            # Step 1c: Insert nodule measurement table.
            # The design places it directly under the schema, so that is the
            # default position; a [TABLEAU] placeholder overrides it.
            if nodule_table and nodule_table.get("rows"):
                anchor, where = None, ""
                for paragraph in doc.paragraphs:
                    if "[TABLEAU]" in paragraph.text:
                        for run in paragraph.runs:
                            run.text = ""
                        anchor, where = paragraph, "[TABLEAU] placeholder"
                        break

                if anchor is None and schema_paragraph is not None:
                    anchor, where = schema_paragraph, "under the schema"
                if anchor is None:
                    anchor, where = doc.add_paragraph(), "end of document"

                self._insert_nodule_table(doc, anchor, nodule_table)
                if logger:
                    logger.info(f"Nodule table inserted ({where})")

            # Step 2: Add images at the end
            if image_paths:
                doc.add_page_break()
                for img_path in image_paths:
                    try:
                        if Path(img_path).exists():
                            doc.add_picture(img_path, width=Inches(5.5))
                            doc.add_paragraph("")  # Spacing between images
                        else:
                            if logger:
                                logger.warning(f"Image not found: {img_path}")
                    except Exception as e:
                        if logger:
                            logger.warning(f"Failed to add image {img_path}: {e}")

            # Step 3: Save
            doc.save(output_path)

            if logger:
                logger.info(f"Word document generated: {output_path}")

            return True

        except Exception as e:
            if logger:
                logger.error(f"Failed to generate Word document: {e}", exc_info=e)
            return False

    # ------------------------------------------------------------------
    # Tableau de mesures des nodules (design handoff schema_thyroidien)
    # ------------------------------------------------------------------

    # Colonnes du design : les trois dimensions et le volume sont alignes a droite
    _TABLE_KEYS = ["nodule", "cote", "siege", "long_mm", "larg_mm",
                   "epais_mm", "volume_ml", "examen"]
    _TABLE_NUMERIC = {"long_mm", "larg_mm", "epais_mm", "volume_ml"}
    # Largeurs en cm, total 16 cm = largeur utile d'une page A4. « Siège » est
    # la seule colonne a texte libre, elle prend la place que les autres
    # n'utilisent pas.
    _TABLE_WIDTHS_CM = [1.4, 1.5, 4.0, 1.75, 1.75, 1.8, 1.9, 1.9]

    def _insert_nodule_table(self, doc, anchor_paragraph, table_data: dict):
        """Insert the nodule measurement table right after `anchor_paragraph`."""
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        headers = table_data["headers"]
        rows = table_data["rows"]

        # 1 en-tete + n nodules + 1 ligne de total
        table = doc.add_table(rows=len(rows) + 2, cols=len(headers))

        # Word ignore la largeur d'une colonne si elle n'est pas repetee sur
        # chaque cellule (tcW) ; la grille (gridCol) seule ne suffit pas, et
        # l'autofit recalculerait tout de toute facon.
        table.autofit = False
        for col, width_cm in zip(table.columns, self._TABLE_WIDTHS_CM):
            col.width = Cm(width_cm)
            for cell in col.cells:
                cell.width = Cm(width_cm)

        header_color = RGBColor(0x6B, 0x68, 0x64)
        body_color = RGBColor(0x23, 0x23, 0x23)

        def write(cell, text, *, size, color, bold=False, right=False):
            para = cell.paragraphs[0]
            # Vider la cellule sans laisser de run vide non formate
            for existing in list(para.runs):
                existing._element.getparent().remove(existing._element)
            if right:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(text)
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold

        # En-tete : 11 px capitales, filet fonce dessous
        for col, label in enumerate(headers):
            write(table.cell(0, col), label.upper(), size=8.5,
                  color=header_color, bold=True,
                  right=self._TABLE_KEYS[col] in self._TABLE_NUMERIC)
        self._set_row_bottom_border(table.rows[0], "232323", 12)

        # Corps : 12 px, filet clair sous chaque ligne
        for i, row_data in enumerate(rows, start=1):
            for col, key in enumerate(self._TABLE_KEYS):
                write(table.cell(i, col), str(row_data.get(key, "")), size=9,
                      color=body_color, right=key in self._TABLE_NUMERIC)
            border = "232323" if i == len(rows) else "ECE9E4"
            self._set_row_bottom_border(table.rows[i], border,
                                        12 if i == len(rows) else 4)

        # Ligne de total : libelle sur les colonnes de gauche, volume aligne
        total_row = table.rows[-1]
        write(total_row.cells[0], table_data["total_label"], size=9,
              color=body_color, bold=True)
        for col in range(1, len(headers)):
            key = self._TABLE_KEYS[col]
            text = table_data["total_volume_ml"] if key == "volume_ml" else ""
            write(total_row.cells[col], text, size=9, color=body_color,
                  bold=(key == "volume_ml"), right=key in self._TABLE_NUMERIC)

        # python-docx ajoute la table en fin de document : la remonter au placeholder
        anchor_paragraph._p.addnext(table._tbl)

    @staticmethod
    def _set_row_bottom_border(row, hex_color: str, size: int):
        """Set a bottom border on every cell of `row` (size in eighths of a point)."""
        from docx.oxml.ns import qn

        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn('w:tcBorders'))
            if borders is None:
                borders = tc_pr.makeelement(qn('w:tcBorders'), {})
                tc_pr.append(borders)
            bottom = borders.find(qn('w:bottom'))
            if bottom is None:
                bottom = borders.makeelement(qn('w:bottom'), {})
                borders.append(bottom)
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), str(size))
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), hex_color)

    def _format_measurements(self, measurements: List) -> str:
        """Format measurements into medical report text"""
        # Separate measurements by type
        right_lobe = next((m for m in measurements
                          if m.side == "RT" and not m.nodule and not m.is_isthmus), None)
        left_lobe = next((m for m in measurements
                         if m.side == "LT" and not m.nodule and not m.is_isthmus), None)
        # « ISTHME » seul = mesure de la glande ; « ISTHME N2 » = nodule situe
        # dans l'isthme, qui ne doit pas etre pris pour la mesure de glande.
        isthmus = next((m for m in measurements
                        if m.is_isthmus and not m.nodule), None)
        # Tries par numero : l'ordre d'extraction n'est pas celui des numeros
        nodules = sorted(
            (m for m in measurements if m.nodule),
            key=lambda m: int(m.nodule) if m.nodule.isdigit() else 999
        )

        # Build report text (using \r\n for Word compatibility)
        text = "• Volume thyroïdien\r\n"
        text += f"- lobe droit : {right_lobe.text if right_lobe else 'non mesuré'}\r\n"
        text += f"- lobe gauche : {left_lobe.text if left_lobe else 'non mesuré'}\r\n"
        text += f"- isthme : {isthmus.text if isthmus else 'non mesuré'}\r\n"
        text += "• Echogénicité glandulaire homogène\r\n"
        text += "• Pas d'anomalie de la vascularisation\r\n"
        text += "• Nodules :\r\n"

        # Meme forme que ThyroidReport.get_formatted_text (mode SR/hybride),
        # pour que le compte rendu soit identique quel que soit le mode.
        for nodule in nodules:
            if nodule.is_isthmus:
                location = "isthme"
            elif nodule.side == "RT":
                location = "lobe droit"
            else:
                location = "lobe gauche"
            text += f"  - Nodule {nodule.nodule} ({location}) : {nodule.text}\r\n"

        text += "• Etude des ganglions (secteurs II, III, IV, VI) et du tractus thyréoglosse : 0"

        return text

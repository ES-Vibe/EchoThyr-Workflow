"""
Word document generation module using python-docx
Generates medical reports from template with bookmarks
"""

from docx import Document
from docx.shared import Inches
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass


@dataclass
class PatientInfo:
    """Patient information"""
    last_name: str = "A PRECISER"
    first_name: str = ""
    exam_date: str = ""


class WordGenerator:
    """Generate Word documents from template"""

    def __init__(self, template_path: str):
        self.template_path = template_path

    def extract_patient_info(self, folder_name: str) -> PatientInfo:
        """Extract patient info from folder name (format: 'NOM Prenom')"""
        from datetime import datetime

        parts = folder_name.split()
        info = PatientInfo(exam_date=datetime.now().strftime("%d.%M.%Y"))

        if len(parts) >= 1:
            info.last_name = parts[0].upper()
        if len(parts) >= 2:
            info.first_name = parts[1]

        return info

    def generate_report(
        self,
        patient_info: PatientInfo,
        measurements: List,
        image_paths: List[str],
        output_path: str,
        logger=None
    ) -> bool:
        """
        Generate medical report from template

        Args:
            patient_info: Patient information
            measurements: List of Measurement objects
            image_paths: List of resized image paths
            output_path: Path for output DOCX file
            logger: Optional logger

        Returns:
            True if successful, False otherwise
        """
        try:
            # Load template
            doc = Document(self.template_path)

            # Replace bookmarks/placeholders
            self._replace_text(doc, "«NOM»", patient_info.last_name)
            self._replace_text(doc, "«PRENOM»", patient_info.first_name)
            self._replace_text(doc, "«DATE»", patient_info.exam_date)

            # Generate measurement text
            measurement_text = self._format_measurements(measurements)
            self._replace_text(doc, "«RESULTAT»", measurement_text)

            # Add images at the end
            if image_paths:
                doc.add_page_break()
                for img_path in image_paths:
                    try:
                        doc.add_picture(img_path, width=Inches(6))
                        doc.add_paragraph()
                    except Exception as e:
                        if logger:
                            logger.warning(f"Failed to add image {img_path}: {e}")

            # Save document
            doc.save(output_path)

            if logger:
                logger.success(f"Word document generated: {output_path}")

            return True

        except Exception as e:
            if logger:
                logger.error(f"Failed to generate Word document: {e}", exc_info=e)
            return False

    def _replace_text(self, doc: Document, placeholder: str, replacement: str):
        """Replace placeholder text in document"""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, replacement)

    def _format_measurements(self, measurements: List) -> str:
        """Format measurements into medical report text"""
        # Separate measurements by type
        right_lobe = next((m for m in measurements
                          if m.side == "RT" and not m.nodule and not m.is_isthmus), None)
        left_lobe = next((m for m in measurements
                         if m.side == "LT" and not m.nodule and not m.is_isthmus), None)
        isthmus = next((m for m in measurements if m.is_isthmus), None)
        nodules = [m for m in measurements if m.nodule]

        # Build report text
        text = "• Volume thyroïdien\n"
        text += f"- lobe droit : {right_lobe.text if right_lobe else 'non mesuré'}\n"
        text += f"- lobe gauche : {left_lobe.text if left_lobe else 'non mesuré'}\n"
        text += f"- isthme : {isthmus.text if isthmus else 'non mesuré'}\n"
        text += "• Echogénicité glandulaire homogène\n"
        text += "• Pas d'anomalie de la vascularisation\n"
        text += "• Nodules :\n"

        for nodule in nodules:
            location = "Lobe droit" if nodule.side == "RT" else "Lobe gauche"
            text += f"  - Nodule N{nodule.nodule} {location} : {nodule.text}\n"

        text += "• Etude des ganglions (secteurs II, III, IV, VI) et du tractus thyréoglosse : 0"

        return text
